import re
from pathlib import Path
from PyPDF2 import PdfReader
from difflib import SequenceMatcher
from io import BytesIO


def extract_text_from_pdf(file_obj) -> str:
    """Extraer texto de un archivo PDF (acepta path string o file object)"""
    try:
        # Si es un string (path), abrirlo
        if isinstance(file_obj, str):
            with open(file_obj, "rb") as f:
                reader = PdfReader(f)
                return "\n".join((p.extract_text() or "") for p in reader.pages)
        # Si es un file object (BytesIO o similar)
        else:
            reader = PdfReader(file_obj)
            return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as e:
        print(f"❌ Error extrayendo texto de PDF: {e}")
        return ""

def extract_text_from_docx(file_obj) -> str:
    """Extraer texto de un archivo DOCX (acepta path string o file object)"""
    try:
        from docx import Document
        # Si es un string (path), abrirlo normalmente
        if isinstance(file_obj, str):
            doc = Document(file_obj)
        # Si es un file object
        else:
            doc = Document(file_obj)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"❌ Error extrayendo texto de DOCX: {e}")
        return ""

def extract_text(file_field) -> str:
    """
    Extraer texto de un archivo. 
    Acepta:
    - String path (almacenamiento local)
    - Django FileField (Cloudinary u otro storage)
    """
    try:
        # Si es un string, usar el método antiguo
        if isinstance(file_field, str):
            ext = Path(file_field).suffix.lower()
            if ext == ".pdf":
                return extract_text_from_pdf(file_field)
            if ext == ".docx":
                return extract_text_from_docx(file_field)
            return ""
        
        # Si es un FileField (Cloudinary), leer el contenido
        # Obtener la extensión del nombre del archivo
        filename = file_field.name
        ext = Path(filename).suffix.lower()
        
        print(f"📖 Leyendo archivo desde storage: {filename}")
        
        # Asegurarse de que el puntero del archivo esté al inicio
        try:
            file_field.seek(0)
        except:
            pass
        
        # Leer el contenido del archivo como BytesIO
        file_content = BytesIO(file_field.read())
        
        print(f"📖 Archivo leído: {file_content.getbuffer().nbytes} bytes")
        
        if ext == ".pdf":
            return extract_text_from_pdf(file_content)
        if ext == ".docx":
            return extract_text_from_docx(file_content)
        
        return ""
    except Exception as e:
        print(f"❌ Error en extract_text: {e}")
        import traceback
        traceback.print_exc()
        return ""

def compute_match(required_skills, candidate_skills):
    if not required_skills or not candidate_skills:
        return 0.0

    req = {s.lower() for s in required_skills}
    cand = {s.lower() for s in candidate_skills}

    intersection = req.intersection(cand)

    score = (len(intersection) / len(req)) * 100
    return round(score, 1)



SKILL_SYNONYMS = {
    # genéricos
    "react.js": "react",
    "react js": "react",
    "reactjs": "react",
    "nest js": "nestjs",
    "nest.js": "nestjs",
    "tsql": "sql",
    "t-sql": "sql",
    "sql server": "sql",
    "mssql": "sql",

    # .NET / C#
    "c sharp": "c#",
    "csharp": "c#",
    ".net": ".net",
    "dotnet": ".net",
    "net": ".net",

    # MAUI
    "maui": "maui",
    ".net maui": "maui",
    ".net con maui": "maui",
}

def normalize_skill(s: str) -> str:
    """
    Normaliza un nombre de skill: minúsculas, sin caracteres raros,
    aplica sinónimos y reconocimiento parcial para MAUI / .NET / C#.
    """
    if not s:
        return ""

    x = s.strip().lower()

    # quitar paréntesis y su contenido
    x = re.sub(r"\(.*?\)", "", x)

    # dejar solo letras, números, espacios, punto y #
    x = re.sub(r"[^\w\s\.#]", " ", x)

    # colapsar espacios
    x = re.sub(r"\s+", " ", x).strip()

    # sinónimos exactos
    if x in SKILL_SYNONYMS:
        return SKILL_SYNONYMS[x]

    # reconocimiento parcial
    if "maui" in x:
        return "maui"

    if ".net" in x:
        return ".net"

    if "c#" in x or "c sharp" in x:
        return "c#"

    # recortes simples: 'framework', 'developer', etc.
    for suffix in [" framework", " developer", " dev", " engineer"]:
        if x.endswith(suffix):
            x = x[: -len(suffix)]

    return x.strip()


def similarity(a: str, b: str) -> float:
    """
    Similaridad de 0 a 1 usando SequenceMatcher de la librería estándar.
    """
    return SequenceMatcher(None, a, b).ratio()


def compute_match_v2(required_skills, candidate_skills):
    """
    Calcula un match score mejorado:
      - Normaliza skills (minúsculas + sinónimos).
      - Usa similitud difusa (fuzzy) para emparejar.
      - Da un bonus pequeño por amplitud de stack.
    """
    if not required_skills:
        return 0.0

    # Normalizar listas
    req_norm = [normalize_skill(s) for s in required_skills if s]
    cand_norm = [normalize_skill(s) for s in candidate_skills or [] if s]

    # eliminar vacíos
    req_norm = [s for s in req_norm if s]
    cand_norm = [s for s in cand_norm if s]

    if not req_norm or not cand_norm:
        return 0.0

    per_skill_scores = []

    for req in req_norm:
        best = 0.0
        for cand in cand_norm:
            sim = similarity(req, cand)

            if req and cand and (req in cand or cand in req):
                sim = max(sim, 0.95)

            if sim > best:
                best = sim

        if best >= 0.90:
            weight = 1.0      # match casi exacto o contenido dentro
        elif best >= 0.75:
            weight = 0.7
        elif best >= 0.60:
            weight = 0.4
        else:
            weight = 0.0

        per_skill_scores.append(weight)

    # Score base: promedio de pesos (0–1) → 0–100
    if per_skill_scores:
        base_score = (sum(per_skill_scores) / len(per_skill_scores)) * 100.0
    else:
        base_score = 0.0

    # BONUS por amplitud de stack (máx 10 puntos extra)
    unique_stack = set(cand_norm)
    breadth_bonus = min(len(unique_stack) / 10.0, 1.0) * 10.0

    final_score = base_score * 0.9 + breadth_bonus * 0.1  # 90% skills requeridos, 10% amplitud
    return round(min(final_score, 100.0), 1)